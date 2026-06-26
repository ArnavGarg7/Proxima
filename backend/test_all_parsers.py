import os
import asyncio
import docx
from fpdf import FPDF
from proxima.database import AsyncSessionLocal
from proxima.services.chunking import ChunkingService
from proxima.services.document_ingestion import DocumentIngestionService
from proxima.models.core import Document, User
from sqlalchemy import select

def create_dummy_files():
    # 1. TXT
    with open("test_doc.txt", "w", encoding="utf-8") as f:
        f.write("This is a simple text file.\n\nIt has multiple paragraphs.\n   Some extra spaces here.")

    # 2. DOCX
    doc = docx.Document()
    doc.add_heading('DOCX Engineering Notes', 0)
    doc.add_paragraph('This is a test paragraph for the DOCX format.\n\n\nIt should be normalized.')
    doc.save("test_doc.docx")

    # 3. PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="PDF Engineering Notes", ln=1, align="C")
    pdf.cell(200, 10, txt="This is a test paragraph for the PDF format.", ln=2, align="L")
    pdf.output("test_doc.pdf")
    
    # 4. Bad PDF (simulate no extractable text by making an empty pdf or invalid one)
    # We'll just create an empty pdf
    pdf_empty = FPDF()
    pdf_empty.add_page()
    pdf_empty.output("test_empty.pdf")

async def test_ingestion(file_path: str):
    print(f"\n--- Testing {file_path} ---")
    async with AsyncSessionLocal() as db:
        user = (await db.execute(select(User).limit(1))).scalar_one_or_none()
        if not user:
            user = User(email="test@example.com", name="Test User")
            db.add(user)
            await db.commit()
            
        new_doc = Document(user_id=user.user_id, title=file_path, status="uploaded")
        db.add(new_doc)
        await db.commit()
        
        chunking_svc = ChunkingService()
        ingestion_svc = DocumentIngestionService(chunking_svc, db)
        
        await ingestion_svc.ingest_document(str(new_doc.document_id), file_path)
        
        from proxima.models.core import DocumentChunk
        chunks = (await db.execute(select(DocumentChunk).where(DocumentChunk.document_id == new_doc.document_id).order_by(DocumentChunk.chunk_index))).scalars().all()
        
        await db.refresh(new_doc)
        print(f"Status: {new_doc.status}")
        print(f"Chunks created: {len(chunks)}")
        if chunks:
            print(f"First chunk preview: {chunks[0].content[:100]!r}")

async def run_all():
    create_dummy_files()
    await test_ingestion("test_doc.txt")
    await test_ingestion("test_doc.docx")
    await test_ingestion("test_doc.pdf")
    await test_ingestion("test_empty.pdf")

if __name__ == "__main__":
    asyncio.run(run_all())
