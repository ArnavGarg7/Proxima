import os
import time
import requests
import docx
from uuid import uuid4

def test_docx_upload():
    # 1. Create a dummy docx file
    doc = docx.Document()
    doc.add_heading('AI Engineering Notes', 0)
    doc.add_paragraph('This is a test paragraph for the AI engineering notes document.')
    doc.add_paragraph('It contains some details about the Proxima architecture.')
    
    file_path = "test_ai_engineering.docx"
    doc.save(file_path)
    
    print(f"Created dummy file: {file_path}")
    
    # We don't have a token, but the user is logged in via browser.
    # For testing, we can just insert it into DB directly to verify extraction
    # Wait, the ingestion service is just a function. Let's call it directly!
    import asyncio
    from proxima.database import AsyncSessionLocal
    from proxima.services.chunking import ChunkingService
    from proxima.services.document_ingestion import DocumentIngestionService
    from proxima.models.core import Document, User
    
    async def run_test():
        async with AsyncSessionLocal() as db:
            # Create dummy user if not exists
            from sqlalchemy import select
            user = (await db.execute(select(User).limit(1))).scalar_one_or_none()
            if not user:
                user = User(email="test@example.com", name="Test User")
                db.add(user)
                await db.commit()
                
            new_doc = Document(user_id=user.user_id, title="test_ai_engineering.docx", status="uploaded")
            db.add(new_doc)
            await db.commit()
            
            print(f"Created Document in DB: {new_doc.document_id}")
            
            chunking_svc = ChunkingService()
            ingestion_svc = DocumentIngestionService(chunking_svc, db)
            
            print("Running ingestion...")
            success = await ingestion_svc.ingest_document(str(new_doc.document_id), file_path)
            print(f"Ingestion success: {success}")
            
            # Check chunks
            from proxima.models.core import DocumentChunk
            chunks = (await db.execute(select(DocumentChunk).where(DocumentChunk.document_id == new_doc.document_id))).scalars().all()
            
            print(f"Created {len(chunks)} chunks.")
            for c in chunks:
                print(f"Chunk {c.chunk_index}: {c.content}")
                
            await db.refresh(new_doc)
            print(f"Final Document Status: {new_doc.status}")

    asyncio.run(run_test())

if __name__ == "__main__":
    test_docx_upload()
